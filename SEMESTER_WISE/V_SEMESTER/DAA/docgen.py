from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pdfkit

# Create a new Document
doc = Document()

# Define a style for section headers
def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.space_after = Pt(6)

# Set up the title and contact information
title = doc.add_paragraph()
title_run = title.add_run('VIJAY KUMAR')
title_run.bold = True
title_run.font.size = Pt(20)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

contact_info = doc.add_paragraph()
contact_info.add_run("Sasaram, Bihar 821115 | ").bold = True
contact_info.add_run("+919727573503 | ").bold = True
contact_info.add_run("babauity@gmail.com | ").bold = True
contact_info.add_run("linkedin.com/4u45").bold = True
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a horizontal line
hr = OxmlElement('w:p')
hr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p = OxmlElement('w:p')
r = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
r.append(rPr)
p.append(r)
hr.append(p)
hr.append(OxmlElement('w:br'))
doc.add_paragraph()._element.append(hr)

# Add Objective section
add_section_header('Objective')
doc.add_paragraph("To secure a challenging position in the field of IT where I can utilize my skills in network security, cybersecurity, programming, and IT support to contribute to organizational growth while continuing to learn and develop professionally.")

# Add Education section
add_section_header('Education')

education = [
    ("B.Tech in Computer Science and Engineering", "Parul University, Vadodara, Gujarat", "Expected Graduation: 2026"),
    ("High School (12th Grade)", "Year of Graduation: 2021"),
    ("Secondary School (10th Grade)", "Year of Graduation: 2019")
]

for degree, *details in education:
    p = doc.add_paragraph()
    p.add_run(degree).bold = True
    for detail in details:
        p.add_run(f"\n{detail}")

# Add Skills section
add_section_header('Skills')

skills = [
    "Network Fundamentals (CCNA)",
    "Routing and Switching (CCNA)",
    "Network Security (CCNA)",
    "IP Connectivity (CCNA)",
    "Network Automation and Programmability (CCNA)",
    "Cybersecurity Foundations (Google Cybersecurity Professional)",
    "Incident Response and Handling (Google Cybersecurity Professional)",
    "Security Operations (Google Cybersecurity Professional)",
    "Risk Management (Google Cybersecurity Professional)",
    "IT Support",
    "Proficient in C, Java, Python, MySQL"
]

for skill in skills:
    doc.add_paragraph(skill, style='List Bullet')

# Add Certifications section
add_section_header('Certifications')

certifications = [
    "Diploma in Computer Application (Completed: February 2021)",
    "CCNA Certificate",
    "Google Cybersecurity Professional Certificate",
    "IT Support Specialist",
]

for certification in certifications:
    doc.add_paragraph(certification, style='List Bullet')

# Add Activities and Interests section
add_section_header('Activities and Interests')
doc.add_paragraph("Any relevant activities or interests that demonstrate your skills or character.")

# Add References section if space allows
add_section_header('References')
doc.add_paragraph("Available upon request.")

# Adjust page layout to narrow margins and ensure content fits on one page
sections = doc.sections
for section in sections:
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

# Save the document as HTML
html_path = "resume.html"
doc.save(html_path)

# Convert HTML to PDF
pdf_path = "resume.pdf"
pdfkit.from_file(html_path, pdf_path)

print(f'Resume saved as {pdf_path}')
